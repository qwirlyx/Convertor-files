import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


def clean_text(text: str) -> str:
    return (
        text.replace("\u00ad", "")
        .replace("\u200b", "")
        .replace("\ufeff", "")
        .replace("\ufffe", "")
        .replace("\xa0", " ")
        .strip()
    )


def paragraph_has_drawing(paragraph) -> bool:
    xml = paragraph._element.xml
    return "<w:drawing" in xml or "<pic:pic" in xml or "<w:object" in xml


def remove_paragraph(paragraph) -> None:
    el = paragraph._element
    parent = el.getparent()
    if parent is not None:
        parent.remove(el)


def convert_anchor_to_inline(paragraph) -> bool:
    converted_any = False
    p_el = paragraph._element

    drawings = p_el.xpath(".//w:drawing")
    for drawing in drawings:
        anchors = drawing.xpath("./wp:anchor")
        for anchor in anchors:
            inline = OxmlElement("wp:inline")
            inline.set("distT", "0")
            inline.set("distB", "0")
            inline.set("distL", "0")
            inline.set("distR", "0")

            for xpath_expr in (
                "./wp:extent",
                "./wp:effectExtent",
                "./wp:docPr",
                "./wp:cNvGraphicFramePr",
                "./a:graphic",
            ):
                nodes = anchor.xpath(xpath_expr)
                if nodes:
                    inline.append(deepcopy(nodes[0]))

            parent = anchor.getparent()
            if parent is not None:
                parent.replace(anchor, inline)
                converted_any = True

    return converted_any

def fix_symbol_bullets(doc: Document) -> None:
    bullet_like_chars = {
        "•", "", "·", "▪", "◦", "",
        "□", "■", "▫", "▢", "▣", "◻", "◼", "⬜", "⬛",
        "☐", "☑", "❑", "❒", "◽", "◾", "○", "●"
    }

    dingbat_fonts = (
        "symbol",
        "wingdings",
        "wingdings 2",
        "wingdings 3",
        "webdings",
        "zapfdingbats",
        "dingbats",
    )

    def looks_like_list_marker(run_text: str, font_name: str) -> bool:
        stripped = (run_text or "").strip()
        if not stripped:
            return False

        if stripped in bullet_like_chars:
            return True

        if any(name in font_name for name in dingbat_fonts):
            return True

        if len(stripped) == 1:
            ch = stripped[0]
            # любой одиночный "спецсимвол", который не буква и не цифра
            if not ch.isalnum():
                code = ord(ch)
                # misc symbols / geometric / dingbats / private use
                if (
                    0x2190 <= code <= 0x2BFF
                    or 0xE000 <= code <= 0xF8FF
                ):
                    return True

        return False

    for p in doc.paragraphs:
        runs = list(p.runs)
        if not runs:
            continue

        for i, run in enumerate(runs):
            txt = run.text or ""
            stripped = txt.strip()
            font_name = (run.font.name or "").strip().lower()

            if not looks_like_list_marker(txt, font_name):
                continue

            # не трогаем длинные куски текста
            if len(stripped) > 2:
                continue

            # чтобы не ломать обычные символы в тексте:
            # считаем маркером только run перед табом/текстом списка
            next_text = ""
            if i + 1 < len(runs):
                next_text = runs[i + 1].text or ""

            prev_text = ""
            if i - 1 >= 0:
                prev_text = runs[i - 1].text or ""

            is_list_position = (
                prev_text in {"", "\t", "\n"}
                or txt.startswith("\t")
                or next_text.startswith("\t")
                or (i + 2 < len(runs) and (runs[i + 2].text or "").strip() != "")
            )

            if not is_list_position:
                continue

            run.text = "•"
            run.font.name = "Times New Roman"
            run.bold = True

            rPr = run._element.get_or_add_rPr()
            rFonts = rPr.rFonts
            if rFonts is None:
                rFonts = OxmlElement("w:rFonts")
                rPr.append(rFonts)

            rFonts.set(qn("w:ascii"), "Times New Roman")
            rFonts.set(qn("w:hAnsi"), "Times New Roman")
            rFonts.set(qn("w:cs"), "Times New Roman")

def postprocess_docx(path: Path) -> None:
    doc = Document(str(path))

    fix_symbol_bullets(doc)

    for p in list(doc.paragraphs):
        if paragraph_has_drawing(p):
            converted = convert_anchor_to_inline(p)

            if clean_text(p.text) == "":
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                pf = p.paragraph_format
                pf.space_before = Pt(6)
                pf.space_after = Pt(6)

            if converted:
                # После перевода в inline Word обычно перестает класть картинки
                # поверх текста или под текст.
                pass
            continue

        text = clean_text(p.text)
        if text in {"|", "/", "\\", "_"}:
            remove_paragraph(p)

    doc.save(str(path))


def main():
    if len(sys.argv) < 3:
        print("usage: pdf_to_docx_layout.py input.pdf output.docx", file=sys.stderr)
        sys.exit(2)

    src = Path(sys.argv[1])
    dst = Path(sys.argv[2])

    if not src.is_file():
        print("source pdf not found", file=sys.stderr)
        sys.exit(3)

    try:
        from pdf2docx import Converter
    except Exception as e:
        print(f"pdf2docx import error: {e}", file=sys.stderr)
        sys.exit(4)

    cv = None
    try:
        cv = Converter(str(src))
        cv.convert(str(dst), start=0, end=None)
    except Exception as e:
        print(f"pdf2docx convert error: {e}", file=sys.stderr)
        sys.exit(5)
    finally:
        if cv is not None:
            try:
                cv.close()
            except Exception:
                pass

    if not dst.is_file() or dst.stat().st_size <= 0:
        print("output docx not created", file=sys.stderr)
        sys.exit(6)

    try:
        postprocess_docx(dst)
    except Exception as e:
        print(f"postprocess error: {e}", file=sys.stderr)
        sys.exit(7)


if __name__ == "__main__":
    main()
